import streamlit as st
from groq import Groq
from supabase import create_client, Client
from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from datetime import datetime
import io
import re
import requests
import urllib.parse

# ─────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────
GROQ_API_KEY    = st.secrets["GROQ_API_KEY"]
SUPABASE_URL    = st.secrets["SUPABASE_URL"]
SUPABASE_KEY    = st.secrets["SUPABASE_KEY"]

groq_client = Groq(api_key=GROQ_API_KEY)
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ─────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────
st.set_page_config(page_title="AI Teacher Assistant", page_icon="📚", layout="wide")


# ─────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────
for key, val in {
    "user": None,
    "token": None,
    "refresh_token": None,
    "current_output": "",
    "current_title": "",
    "current_task": "",
    "edit_mode": False,
    "edit_id": None,
    "saved_plans": [],
    "auth_mode": "login",
    "generated_images": {},
}.items():
    if key not in st.session_state:
        st.session_state[key] = val

# Re-attach session on every script rerun (Streamlit creates a fresh
# supabase client each run, so RLS won't know who the user is otherwise)
if st.session_state.user and st.session_state.token and st.session_state.refresh_token:
    try:
        supabase.auth.set_session(st.session_state.token, st.session_state.refresh_token)
    except Exception:
        pass

# ─────────────────────────────────────────
#  GRADE OPTIONS
# ─────────────────────────────────────────
GRADE_OPTIONS = {
    "IB (International Baccalaureate)":  ["Grade 11", "Grade 12"],
    "IGCSE (Cambridge)":                  ["Grade 9", "Grade 10"],
    "A-Level (Cambridge)":                ["Grade 11", "Grade 12"],
    "Cambridge Checkpoint":               ["Grade 5", "Grade 6", "Grade 7", "Grade 8"],
    "Singapore (MOE)":                    ["K1", "K2"] + [f"Grade {i}" for i in range(1, 9)],
    "CBSE (India)":                       [f"Grade {i}" for i in range(1, 13)],
    "American (Common Core)":            [f"Grade {i}" for i in range(1, 13)],
    "Australian (ACARA)":                 [f"Grade {i}" for i in range(1, 13)],
    "British (National Curriculum)":      [f"Year {i}" for i in range(1, 14)],
}

# ─────────────────────────────────────────
#  SUBJECT OPTIONS (per curriculum)
# ─────────────────────────────────────────
SUBJECT_OPTIONS = {
    "IB (International Baccalaureate)":  ["Computer Science", "Mathematics", "Physics", "Chemistry",
                                           "Biology", "English", "Economics", "Business Management",
                                           "Psychology", "History", "Geography"],
    "IGCSE (Cambridge)":                  ["Computer Science", "Mathematics", "Physics", "Chemistry",
                                           "Biology", "English", "Economics", "Business Studies",
                                           "ICT", "History", "Geography"],
    "A-Level (Cambridge)":                ["Computer Science", "Mathematics", "Physics", "Chemistry",
                                           "Biology", "English", "Economics", "Business",
                                           "Psychology", "History", "Geography"],
    "Cambridge Checkpoint":               ["Mathematics", "Science", "English", "ICT", "Computer Science"],
    "Singapore (MOE)":                    ["Mathematics", "Science", "English", "Computer Science",
                                           "Social Studies", "Art"],
    "CBSE (India)":                       ["Mathematics", "Science", "English", "Computer Science",
                                           "Social Science", "Hindi", "Physics", "Chemistry", "Biology"],
    "American (Common Core)":            ["Mathematics", "Science", "English Language Arts",
                                           "Computer Science", "Social Studies"],
    "Australian (ACARA)":                 ["Mathematics", "Science", "English", "Digital Technologies",
                                           "Humanities and Social Sciences"],
    "British (National Curriculum)":      ["Mathematics", "Science", "English", "Computing",
                                           "History", "Geography"],
}

# ─────────────────────────────────────────
#  TOPIC OPTIONS (per subject)
# ─────────────────────────────────────────
TOPIC_OPTIONS = {
    "Computer Science": [
        "Abstraction & Decomposition", "Algorithms & Flowcharts", "Programming Basics (Variables, Data Types)",
        "Sequencing, Selection & Iteration", "Functions & Procedures", "Arrays & Lists",
        "Searching & Sorting Algorithms", "Data Structures (Stacks, Queues, Trees)",
        "Object-Oriented Programming", "Databases & SQL", "Networks & Protocols",
        "Binary, Hex & Number Systems", "Logic Gates & Boolean Algebra", "Cyber Security",
        "Web Development (HTML/CSS/JS)", "Python Programming", "Computational Thinking",
        "Software Development Life Cycle", "Ethics in Computing", "Artificial Intelligence Basics"
    ],
    "Mathematics": [
        "Number & Place Value", "Fractions, Decimals & Percentages", "Algebra & Expressions",
        "Equations & Inequalities", "Geometry & Shapes", "Measurement & Units",
        "Statistics & Data Handling", "Probability", "Ratio & Proportion",
        "Functions & Graphs", "Trigonometry", "Calculus (Differentiation/Integration)",
        "Vectors & Transformations", "Sequences & Series", "Coordinate Geometry"
    ],
    "Physics": [
        "Motion & Forces", "Energy & Work", "Electricity & Circuits", "Waves & Sound",
        "Light & Optics", "Magnetism & Electromagnetism", "Thermal Physics",
        "Nuclear Physics & Radioactivity", "Gravitation", "Pressure & Density"
    ],
    "Chemistry": [
        "Atomic Structure & Periodic Table", "Bonding (Ionic, Covalent, Metallic)",
        "Chemical Reactions & Equations", "Acids, Bases & Salts", "Rates of Reaction",
        "Energetics", "Organic Chemistry", "Electrolysis", "States of Matter", "Stoichiometry"
    ],
    "Biology": [
        "Cell Structure & Function", "Nutrition & Digestion", "Respiration", "Photosynthesis",
        "Transport in Plants & Animals", "Reproduction", "Genetics & Inheritance",
        "Evolution & Natural Selection", "Ecosystems & Ecology", "Human Body Systems"
    ],
    "English": [
        "Reading Comprehension", "Creative Writing", "Grammar & Punctuation",
        "Poetry Analysis", "Persuasive Writing", "Narrative Writing", "Shakespeare",
        "Novel Study", "Speaking & Listening", "Vocabulary Building"
    ],
    "English Language Arts": [
        "Reading Comprehension", "Creative Writing", "Grammar & Punctuation",
        "Poetry Analysis", "Persuasive Writing", "Narrative Writing", "Novel Study",
        "Speaking & Listening", "Vocabulary Building"
    ],
    "Science": [
        "Living Things & Habitats", "Forces & Motion", "Materials & Their Properties",
        "Earth & Space", "Energy", "States of Matter", "Human Body", "Plants",
        "Electricity", "Sound & Light"
    ],
    "Economics": [
        "Demand & Supply", "Market Structures", "Macroeconomics & GDP", "Inflation & Unemployment",
        "International Trade", "Government Intervention", "Microeconomics Basics",
        "Elasticity", "Money & Banking", "Economic Development"
    ],
    "Business Management": [
        "Business Organisation", "Marketing", "Finance & Accounts", "Human Resource Management",
        "Operations Management", "Business Strategy", "Entrepreneurship", "External Environment"
    ],
    "Business Studies": [
        "Business Organisation", "Marketing", "Finance & Accounts", "Human Resource Management",
        "Operations Management", "Business Strategy", "Entrepreneurship", "External Environment"
    ],
    "Business": [
        "Business Organisation", "Marketing", "Finance & Accounts", "Human Resource Management",
        "Operations Management", "Business Strategy", "Entrepreneurship", "External Environment"
    ],
    "Psychology": [
        "Research Methods", "Biological Psychology", "Cognitive Psychology", "Developmental Psychology",
        "Social Psychology", "Abnormal Psychology", "Memory", "Attachment"
    ],
    "History": [
        "World War I & II", "Cold War", "Industrial Revolution", "Ancient Civilizations",
        "Independence Movements", "Civil Rights Movements", "Colonialism", "Revolutions"
    ],
    "Geography": [
        "Climate & Weather", "Population & Migration", "Urbanisation", "Natural Resources",
        "Ecosystems & Biomes", "Rivers & Coasts", "Tectonic Hazards", "Globalisation",
        "Development & Inequality"
    ],
    "ICT": [
        "Hardware & Software", "Word Processing & Spreadsheets", "Databases",
        "Networks & Internet", "Multimedia & Presentations", "E-Safety & Digital Citizenship",
        "Programming Basics", "Data Representation"
    ],
    "Social Studies": [
        "Citizenship & Government", "Geography of Singapore/Region", "History of the Nation",
        "Cultural Diversity", "Economics Basics", "Global Connections"
    ],
    "Social Science": [
        "History", "Geography", "Civics", "Economics", "Disaster Management"
    ],
    "Hindi": [
        "व्याकरण (Grammar)", "गद्य (Prose)", "पद्य (Poetry)", "निबंध लेखन (Essay Writing)",
        "पत्र लेखन (Letter Writing)", "अपठित गद्यांश (Unseen Passage)"
    ],
    "Art": [
        "Drawing & Sketching", "Painting Techniques", "Colour Theory", "Sculpture",
        "Art History & Appreciation", "Digital Art"
    ],
    "Digital Technologies": [
        "Algorithms & Programming", "Data Representation", "Digital Systems",
        "Networks", "Impacts of Technology", "Robotics"
    ],
    "Computing": [
        "Algorithms", "Programming (Python/Scratch)", "Data Representation",
        "Networks", "Computer Systems", "Creating Media", "Online Safety"
    ],
    "Humanities and Social Sciences": [
        "History", "Geography", "Civics & Citizenship", "Economics & Business"
    ],
}

# ─────────────────────────────────────────
#  AUTH FUNCTIONS
# ─────────────────────────────────────────
def sign_up(email, password, name, school):
    try:
        res = supabase.auth.sign_up({"email": email, "password": password})
        if res.user:
            supabase.table("teachers").upsert({
                "id": res.user.id,
                "email": email,
                "name": name,
                "school": school
            }).execute()
            return True, "Account created! You can now log in."
        return False, "Signup failed. Try again."
    except Exception as e:
        return False, str(e)

def sign_in(email, password):
    try:
        res = supabase.auth.sign_in_with_password({"email": email, "password": password})
        if res.user:
            st.session_state.user          = res.user
            st.session_state.token         = res.session.access_token
            st.session_state.refresh_token = res.session.refresh_token
            # Attach session so RLS policies see auth.uid() correctly
            supabase.auth.set_session(res.session.access_token, res.session.refresh_token)
            return True, "Login successful!"
        return False, "Invalid email or password."
    except Exception as e:
        return False, str(e)

def reset_password(email):
    try:
        supabase.auth.reset_password_for_email(email)
        return True, "If an account exists for this email, a password reset link has been sent. Check your inbox (and spam folder)."
    except Exception as e:
        return False, str(e)

def update_password(new_password, access_token, refresh_token):
    try:
        supabase.auth.set_session(access_token, refresh_token)
        supabase.auth.update_user({"password": new_password})
        return True, "Password updated successfully! You can now log in with your new password."
    except Exception as e:
        return False, str(e)

def sign_out():
    supabase.auth.sign_out()
    for key in ["user","token","refresh_token","current_output","current_title",
                "current_task","edit_mode","edit_id","saved_plans"]:
        st.session_state[key] = None if key in ["user","token","refresh_token","edit_id"] else \
                                 [] if key == "saved_plans" else \
                                 False if key == "edit_mode" else ""
    st.rerun()

# ─────────────────────────────────────────
#  DATABASE FUNCTIONS
# ─────────────────────────────────────────
def load_plans():
    try:
        uid = st.session_state.user.id
        res = supabase.table("saved_plans").select("*")\
              .eq("teacher_id", uid).order("created_at", desc=True).execute()
        st.session_state.saved_plans = res.data or []
    except:
        st.session_state.saved_plans = []

def save_plan(title, content, task, curriculum, subject, grade):
    try:
        uid = st.session_state.user.id
        if st.session_state.edit_id:
            supabase.table("saved_plans").update({
                "title": title, "content": content,
                "updated_at": datetime.now().isoformat()
            }).eq("id", st.session_state.edit_id).execute()
            st.success("✅ Updated!")
        else:
            supabase.table("saved_plans").insert({
                "teacher_id": uid, "title": title, "content": content,
                "task": task, "curriculum": curriculum,
                "subject": subject, "grade": grade,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat()
            }).execute()
            st.success("✅ Saved!")
        load_plans()
    except Exception as e:
        st.error(f"Save error: {e}")

def delete_plan(plan_id):
    try:
        supabase.table("saved_plans").delete().eq("id", plan_id).execute()
        load_plans()
        st.rerun()
    except Exception as e:
        st.error(f"Delete error: {e}")

# ─────────────────────────────────────────
#  SYLLABUS FUNCTIONS (custom teacher topics)
# ─────────────────────────────────────────
def get_syllabus(curriculum, grade, subject):
    """Return list of topics the teacher has saved for this curriculum/grade/subject, or None."""
    try:
        uid = st.session_state.user.id
        res = supabase.table("syllabus").select("*")\
              .eq("teacher_id", uid)\
              .eq("curriculum", curriculum)\
              .eq("grade", grade)\
              .eq("subject", subject)\
              .execute()
        if res.data:
            topics_str = res.data[0]["topics"]
            return [t.strip() for t in topics_str.split(",") if t.strip()]
        return None
    except Exception:
        return None

def save_syllabus(curriculum, grade, subject, topics_list):
    try:
        uid = st.session_state.user.id
        topics_str = ", ".join(topics_list)
        supabase.table("syllabus").upsert({
            "teacher_id": uid,
            "curriculum": curriculum,
            "grade": grade,
            "subject": subject,
            "topics": topics_str,
            "updated_at": datetime.now().isoformat()
        }, on_conflict="teacher_id,curriculum,grade,subject").execute()
        return True
    except Exception as e:
        st.error(f"Syllabus save error: {e}")
        return False

# ─────────────────────────────────────────
#  AI FUNCTIONS
# ─────────────────────────────────────────
def generate(prompt):
    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": (
                "You are an expert teacher and curriculum designer. "
                "Always produce well-structured, classroom-ready content "
                "with clear headings, bullet points, and timing where relevant."
            )},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2048, temperature=0.7,
    )
    return response.choices[0].message.content

def build_prompt(task, subject, grade, topics, lessons_per_week,
                 duration, curriculum, activity_duration=30, test_duration=45,
                 question_specs="", student_name="", performance_level="",
                 strengths_input="", improvements_input="", comment_tone="",
                 comment_length="", message_platform="", message_purpose="",
                 message_details="", report_period="", topics_covered_input="",
                 homeroom_input=""):
    ctx = {
        "IB (International Baccalaureate)":  "Use IB terminology (ATL skills, Learner Profile, Command Terms). Include inquiry questions.",
        "IGCSE (Cambridge)":                  "Use Cambridge command terms (describe, explain, evaluate, analyse). Align with Cambridge assessment objectives.",
        "A-Level (Cambridge)":                "Use higher-order thinking and Cambridge A-Level mark scheme style.",
        "Cambridge Checkpoint":               "Age-appropriate language and Cambridge Checkpoint assessment style.",
        "Singapore (MOE)":                    "Follow Singapore MOE curriculum. Emphasise problem solving and 21CC skills.",
        "CBSE (India)":                       "Follow CBSE curriculum and NCERT guidelines. Align with board exam pattern.",
        "American (Common Core)":            "Follow Common Core State Standards. Include standards codes.",
        "Australian (ACARA)":                 "Follow Australian Curriculum (ACARA). Include strand references.",
        "British (National Curriculum)":      "Follow UK National Curriculum. Use Ofsted-friendly lesson structure.",
    }.get(curriculum, "")
    base = (f"Curriculum: {curriculum}\nSubject: {subject}\nGrade: {grade}\n"
            f"Topics: {topics}\nLessons/week: {lessons_per_week}\n"
            f"Lesson duration: {duration} min\n\n{ctx}\n\n")

    # Report Card Comment and Parent Message don't need the topics/curriculum base
    if task == "🗣️ Report Card Comment":
        name_ref = student_name.strip() if student_name.strip() else "the student"
        homeroom_line = f"Homeroom: {homeroom_input.strip()}\n" if homeroom_input.strip() else ""
        topics_line = (f"Topics covered this {report_period}: {topics_covered_input.strip()}\n"
                       if topics_covered_input.strip() else "")
        return (
            f"Write a {report_period} report card comment for {name_ref}, a {grade} student "
            f"studying {subject} ({curriculum}).\n"
            f"{homeroom_line}"
            f"{topics_line}"
            f"Overall performance level: {performance_level}\n"
            f"Strengths to mention: {strengths_input if strengths_input.strip() else 'general positive effort'}\n"
            f"Areas to improve: {improvements_input if improvements_input.strip() else 'general next steps for growth'}\n"
            f"Tone: {comment_tone}\n"
            f"Length: {comment_length}\n\n"
            "Write ONLY the comment text, ready to paste directly into a school report card system. "
            "Do not include a greeting, sign-off, or any explanation — just the comment paragraph itself. "
            "If topics covered were given, naturally reference what was learned this period. "
            "Make it specific, professional, and avoid generic filler phrases."
        )

    if task == "💬 Parent Message":
        return (
            f"Write a {message_platform} message for parents about a {grade} {subject} class "
            f"({curriculum}).\n"
            f"Purpose: {message_purpose}\n"
            f"Details to include: {message_details if message_details.strip() else 'general weekly update'}\n\n"
            "Write ONLY the message text, ready to copy and send directly to parents. "
            "Keep it warm, clear, and appropriately brief for the chosen style. "
            "Do not include placeholder brackets like [date] — write it naturally, "
            "or use generic phrasing like 'this week' instead of exact dates."
        )

    return {
        "📋 Lesson Plan": base + (
            f"Create {lessons_per_week} detailed lesson plans, one per lesson this week. "
            f"Each is {duration} minutes. Label as Lesson 1, Lesson 2, etc. "
            "Each includes: Objectives, Materials, Starter (5 min), Main activities with timings, "
            "Assessment, Plenary, Differentiation, Homework."),
        "🎯 Class Activity": base + (
            f"Design an engaging {activity_duration}-minute classroom activity. "
            "Include: Title, Objective, Step-by-step instructions, Group/individual, "
            "Materials, Expected outcomes, Extension task. " + question_specs),
        "📝 Homework": base + (
            "Create a homework assignment. "
            + (question_specs if question_specs else
               "Include 5-8 questions of varying difficulty (easy/medium/hard). ")
            + "Include instructions, expected time, learning outcome."),
        "📄 Unit Test": base + (
            "Create a full unit test covering ALL topics. "
            + (question_specs if question_specs else
               "Include 5 MCQ, 5 short answer, 1 essay/long answer. ")
            + "Show total marks and include mark scheme at end."),
        "📃 Class Test": base + (
            f"Create a {test_duration}-minute class test. "
            + (question_specs if question_specs else
               "Include 3-5 questions, mix of types. ")
            + "Show total marks and include answer key."),
        "❓ Quiz": base + (
            "Create a quiz. "
            + (question_specs if question_specs else
               "Include 10 questions: MCQ, true/false, fill-in-the-blank. ")
            + "Include answers at the end."),
        "📅 Weekly Schedule": base + (
            f"Create a weekly schedule for {lessons_per_week} lessons/week, "
            f"{duration} min each. Include day/time, topic, description, homework slots."),
        "📊 Marking Rubric": base + (
            "Create a marking rubric with 4-5 criteria, 4 performance levels "
            "(Excellent/Good/Satisfactory/Needs Improvement), descriptors, marks. " + question_specs),
        "📘 Exam Base Question": base + (
    "Create a formal exam-style question paper based on curriculum, grade, subject, and topics.\n\n"

    "STRUCTURE:\n"
    "Section A: Multiple Choice Questions (MCQs)\n"
    "Section B: Short Answer Questions\n"
    "Section C: Long Answer / Structured Questions\n\n"

    "REQUIREMENTS:\n"
    "- Follow real board exam patterns (IB / IGCSE / CBSE / A-Level)\n"
    "- Ensure difficulty progression: easy → medium → hard\n"
    "- Use command terms: define, explain, analyse, evaluate\n"
    "- Include application-based reasoning questions\n"
    "- Ensure clarity and proper exam formatting\n\n"

    "MARK SCHEME:\n"
    "At the end, provide a full marking scheme with step-by-step answers and marking points.\n"
    + question_specs
),
    }.get(task, base + f"Create educational content for {task}. " + question_specs)

# ─────────────────────────────────────────
#  EXPORT FUNCTIONS
# ─────────────────────────────────────────
def export_word(title, content, images=None):
    images = images or {}
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    doc.add_paragraph("")
    for line in content.split("\n"):
        m = re.match(r"\[IMAGE:\s*(.*?)\]", line.strip(), re.IGNORECASE)
        if m:
            desc = m.group(1).strip()
            img_bytes = images.get(desc)
            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))
                    cap = doc.add_paragraph(desc)
                    cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else cap.style
                except Exception:
                    doc.add_paragraph(f"[Image: {desc}]")
            else:
                doc.add_paragraph(f"[Image suggestion: {desc}]")
        elif line.startswith("# "):   doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  doc.add_heading(line[3:], level=2)
        elif line.startswith("### "): doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ","* ")): doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip(): doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def export_excel(title, content):
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Lesson Plan"
    hf = PatternFill("solid", fgColor="2E4057")
    ws.merge_cells("A1:C1"); ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = hf; ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A2:C2")
    ws["A2"] = f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}"
    ws["A2"].alignment = Alignment(horizontal="center")
    for col, h in enumerate(["Content","Details","Notes"], 1):
        c = ws.cell(row=3, column=col, value=h)
        c.font = Font(color="FFFFFF", bold=True); c.fill = hf
        c.alignment = Alignment(horizontal="center")
    row = 4
    for line in content.split("\n"):
        if line.strip():
            ws.cell(row=row, column=1, value=line.strip())
            ws.cell(row=row, column=1).alignment = Alignment(wrap_text=True)
            row += 1
    ws.column_dimensions["A"].width = 60
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 30
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# ─────────────────────────────────────────
#  IMAGE FUNCTIONS
# ─────────────────────────────────────────
def extract_image_prompts(content):
    """Find lines like [IMAGE: description] in the AI output and return list of descriptions."""
    return re.findall(r"\[IMAGE:\s*(.*?)\]", content, re.IGNORECASE)

def generate_image(prompt_text, width=512, height=512):
    """Generate an image via Pollinations.ai (free, no API key needed).
    Returns (image_bytes_or_None, debug_info_string)."""
    try:
        encoded = urllib.parse.quote(prompt_text)
        url = (f"https://image.pollinations.ai/prompt/{encoded}"
               f"?width={width}&height={height}&nologo=true&model=flux")
        headers = {"User-Agent": "Mozilla/5.0 (compatible; AITeacherApp/1.0)"}
        resp = requests.get(url, headers=headers, timeout=90)
        ctype = resp.headers.get("content-type", "")
        if resp.status_code == 200 and ctype.startswith("image"):
            return resp.content, "ok"
        else:
            return None, f"status={resp.status_code}, type={ctype}, body={resp.text[:150]}"
    except Exception as e:
        return None, f"exception: {e}"

def reference_image_search_url(query):
    """Build a free reference image search link (Option C)."""
    return f"https://www.google.com/search?q={urllib.parse.quote(query)}&tbm=isch"

# ═════════════════════════════════════════
#  PASSWORD RECOVERY HANDLING
# ═════════════════════════════════════════
import streamlit.components.v1 as components

# Supabase puts the recovery token in the URL *fragment* (#access_token=...),
# which Python can't read directly. Streamlit's components.html iframe is
# sandboxed and browsers BLOCK iframes from calling location.replace() on
# the parent window (confirmed: "current window does not have permission
# to navigate the target frame"). Browsers DO allow a real user click on an
# <a target="_top"> link to navigate the top frame, so instead of forcing
# a JS redirect, we detect the token and show a clickable link/button.
components.html("""
<div id="recovery-link-holder"></div>
<script>
(function() {
    const holder = document.getElementById("recovery-link-holder");
    try {
        const topHash = window.top.location.hash;
        if (topHash && topHash.includes("access_token")) {
            const hash = topHash.substring(1);
            const params = new URLSearchParams(hash);
            const accessToken = params.get("access_token");
            const refreshToken = params.get("refresh_token");
            const type = params.get("type");
            if (accessToken && type === "recovery" &&
                !window.top.location.search.includes("access_token")) {
                const newUrl = window.top.location.pathname +
                    "?access_token=" + encodeURIComponent(accessToken) +
                    "&refresh_token=" + encodeURIComponent(refreshToken) +
                    "&type=" + encodeURIComponent(type);
                const a = document.createElement("a");
                a.href = newUrl;
                a.target = "_blank";
                a.rel = "noopener";
                a.innerText = "🔑 Click here to set your new password →";
                a.style = "display:inline-block;padding:12px 22px;background:#FF4B4B;" +
                          "color:white;font-family:sans-serif;font-size:16px;" +
                          "border-radius:8px;text-decoration:none;font-weight:600;" +
                          "box-shadow:0 2px 6px rgba(0,0,0,0.15);";
                holder.appendChild(a);
                const hint = document.createElement("div");
                hint.innerText = "Opens in a new tab — you can close this old tab afterward.";
                hint.style = "font-family:sans-serif;font-size:12px;color:#888;margin-top:8px;";
                holder.appendChild(hint);
            }
        }
    } catch (e) {
        holder.innerText = "Could not read reset link: " + e.message;
    }
})();
</script>
""", height=90)

qp = st.query_params
is_recovery = qp.get("type") == "recovery" and qp.get("access_token")

if is_recovery and not st.session_state.user:
    st.title("📚 AI Teacher Assistant")
    st.subheader("🔑 Set a new password")
    st.caption("Enter and confirm your new password below.")

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        new_pw  = st.text_input("New Password", type="password", help="Minimum 6 characters")
        new_pw2 = st.text_input("Confirm New Password", type="password")
        if st.button("Update Password", type="primary", use_container_width=True):
            if not new_pw or not new_pw2:
                st.warning("Please fill in both fields.")
            elif new_pw != new_pw2:
                st.error("Passwords do not match.")
            elif len(new_pw) < 6:
                st.error("Password must be at least 6 characters.")
            else:
                ok, msg = update_password(new_pw, qp.get("access_token"), qp.get("refresh_token"))
                if ok:
                    st.success(msg)
                    st.info("Go to the Login tab by refreshing this page, then sign in with your new password.")
                    st.query_params.clear()
                else:
                    st.error(msg)
    st.stop()


# ═════════════════════════════════════════
#  LOGIN / SIGNUP PAGE
# ═════════════════════════════════════════
if not st.session_state.user:
    st.title("📚 AI Teacher Assistant")
    st.caption("Smart lesson planning powered by AI — Free for teachers")
    st.divider()

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        mode = st.radio("", ["🔑 Login", "📝 Sign Up", "❓ Forgot Password"],
                        horizontal=True, label_visibility="collapsed")

        if mode == "🔑 Login":
            st.subheader("Welcome back!")
            email    = st.text_input("Email", placeholder="teacher@school.com")
            password = st.text_input("Password", type="password")
            if st.button("Login", type="primary", use_container_width=True):
                if email and password:
                    ok, msg = sign_in(email, password)
                    if ok:
                        load_plans()
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.warning("Please enter email and password.")
            st.caption("Forgot your password? Click '❓ Forgot Password' above.")

        elif mode == "📝 Sign Up":
            st.subheader("Create your account")
            name     = st.text_input("Your Name",    placeholder="e.g. Mr. Akhilesh")
            school   = st.text_input("School Name",  placeholder="e.g. Singapore International School")
            email    = st.text_input("Email",         placeholder="teacher@school.com")
            password = st.text_input("Password",      type="password", help="Minimum 6 characters")
            password2= st.text_input("Confirm Password", type="password")
            if st.button("Create Account", type="primary", use_container_width=True):
                if not all([name, school, email, password, password2]):
                    st.warning("Please fill in all fields.")
                elif password != password2:
                    st.error("Passwords do not match.")
                elif len(password) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    ok, msg = sign_up(email, password, name, school)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

        else:  # Forgot Password
            st.subheader("Reset your password")
            st.caption("Enter your email and we'll send you a password reset link.")
            reset_email = st.text_input("Email", placeholder="teacher@school.com", key="reset_email")
            if st.button("Send Reset Link", type="primary", use_container_width=True):
                if reset_email:
                    ok, msg = reset_password(reset_email)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
                else:
                    st.warning("Please enter your email.")

    st.stop()

# ═════════════════════════════════════════
#  MAIN APP (logged in)
# ═════════════════════════════════════════
st.title("📚 AI Teacher Assistant")
st.caption("Developed by Akhilesh Kumar Srivastava — AI Teaching Assistant for IB, IGCSE, Singapore, A-Level & more")

with st.expander("ℹ️ About this tool / How to use"):
    st.markdown("""
    **Purpose:** Save time on lesson planning, activities, homework, tests, and quizzes — all generated instantly by AI.

    **Who can use it:** Any teacher — just sign up with your email to get your own private workspace.

    **How to use:**
    1. Select your Curriculum, Grade, Subject, and Topics in the sidebar
    2. Set lessons per week and lesson duration
    3. Choose what to generate (Lesson Plan, Quiz, Test, etc.)
    4. Click ⚡ Generate
    5. Save to Cloud, Edit, or Download as Word / Excel

    **Developed by:** Akhilesh Kumar Srivastava
    """)

# ── Sidebar ───────────────────────────────
with st.sidebar:
    # Teacher info
    st.success(f"👋 {st.session_state.user.email}")
    if st.button("🚪 Logout", use_container_width=True):
        sign_out()

    st.divider()
    st.header("🛠 What do you want to generate?")
    task = st.selectbox("Choose a task", [
        "📋 Lesson Plan", "🎯 Class Activity", "📝 Homework",
        "📄 Unit Test", "📃 Class Test", "❓ Quiz",
        "📅 Weekly Schedule", "📊 Marking Rubric",
        "🗣️ Report Card Comment", "💬 Parent Message", "📘 Exam Base Question",
    ], label_visibility="collapsed")

    needs_curriculum_fields = task not in ["🗣️ Report Card Comment", "💬 Parent Message"]

    st.divider()
    st.header("⚙️ Settings")

    if needs_curriculum_fields:
        curriculum = st.selectbox("🎓 Curriculum", list(GRADE_OPTIONS.keys()))
        grade      = st.selectbox("📊 Grade / Year", GRADE_OPTIONS[curriculum])

        subject_list = SUBJECT_OPTIONS.get(curriculum, ["Computer Science", "Mathematics", "Science", "English"])
        subject_list = subject_list + ["Other (type manually)"]
        subject_choice = st.selectbox("📖 Subject", subject_list)

        if subject_choice == "Other (type manually)":
            subject = st.text_input("Enter subject name", placeholder="e.g. Computer Science")
        else:
            subject = subject_choice

        # ── Topics: use custom syllabus if teacher has saved one, else built-in list ──
        custom_topics = get_syllabus(curriculum, grade, subject)
        default_list  = TOPIC_OPTIONS.get(subject, [])
        topic_list    = custom_topics if custom_topics is not None else default_list

        if topic_list:
            selected_topics = st.multiselect(
                "📌 Topics (select one or more)",
                topic_list,
                help="Select topics to include — you can also add extra topics below"
            )
        else:
            selected_topics = []
            st.info("📋 No topic list yet for this subject/grade. Add your syllabus below, or type topics manually.")

        extra_topics = st.text_area(
            "➕ Additional topics (optional)",
            placeholder="Type any extra topics not in the list above, separated by commas",
            height=60
        )

        topic_parts = selected_topics + (
            [t.strip() for t in extra_topics.split(",") if t.strip()] if extra_topics else []
        )
        topics = ", ".join(topic_parts)

        # ── Syllabus upload / edit ──
        with st.expander("📚 My Syllabus (add / edit topics for this subject)"):
            st.caption(f"For: {curriculum} · {grade} · {subject}")
            if custom_topics:
                st.success(f"✅ You have {len(custom_topics)} custom topics saved for this subject.")
                existing_text = ", ".join(custom_topics)
            else:
                st.caption("No custom syllabus saved yet. Add your topics below (comma separated) — "
                           "these will appear in the dropdown above for future use.")
                existing_text = ", ".join(default_list)

            syllabus_text = st.text_area(
                "Topics (comma separated)",
                value=existing_text,
                height=120,
                key=f"syllabus_{curriculum}_{grade}_{subject}"
            )
            if st.button("💾 Save Syllabus", use_container_width=True,
                          key=f"save_syllabus_{curriculum}_{grade}_{subject}"):
                new_topics = [t.strip() for t in syllabus_text.split(",") if t.strip()]
                if new_topics:
                    if save_syllabus(curriculum, grade, subject, new_topics):
                        st.success("✅ Syllabus saved! It will now appear in the Topics dropdown.")
                        st.rerun()
                else:
                    st.warning("Please enter at least one topic.")

        lessons_per_week = st.number_input("Lessons per week", min_value=1, max_value=10, value=3)
        duration         = st.number_input("Lesson duration (min)", min_value=20, max_value=180, value=45)

    else:
        # Report Card Comment / Parent Message still need a light Subject + Grade
        # context, but skip the full curriculum/topics/syllabus machinery.
        curriculum = st.selectbox("🎓 Curriculum", list(GRADE_OPTIONS.keys()))
        grade      = st.selectbox("📊 Grade / Year", GRADE_OPTIONS[curriculum])
        subject_list = SUBJECT_OPTIONS.get(curriculum, ["Computer Science", "Mathematics", "Science", "English"])
        subject_list = subject_list + ["Other (type manually)"]
        subject_choice = st.selectbox("📖 Subject", subject_list)
        if subject_choice == "Other (type manually)":
            subject = st.text_input("Enter subject name", placeholder="e.g. Computer Science")
        else:
            subject = subject_choice
        topics = ""
        lessons_per_week = 1
        duration = 45

    st.divider()

    activity_duration = duration
    test_duration     = 45

    if task == "🎯 Class Activity":
        st.info("⏱ Activity duration?")
        activity_duration = st.number_input("Activity duration (min)", min_value=5, max_value=180, value=30)
    elif task == "📃 Class Test":
        st.info("⏱ Class test duration?")
        test_duration = st.number_input("Test duration (min)", min_value=10, max_value=180, value=45)
    elif task == "📋 Lesson Plan":
        st.info(f"📅 Will generate {int(lessons_per_week)} lesson plan(s)")

    # ── Report Card Comment specific inputs ──
    student_name        = ""
    performance_level   = ""
    strengths_input      = ""
    improvements_input   = ""
    comment_tone         = "Warm and encouraging"
    comment_length       = "Medium (3-4 sentences)"
    report_mode          = "Single student"
    batch_groups         = []   # list of (level, [names])
    report_period        = "Term 1"
    topics_covered_input = ""
    homeroom_input       = ""

    if task == "🗣️ Report Card Comment":
        report_mode = st.radio(
            "How many students?",
            ["Single student", "Whole class (batch)"],
            horizontal=True
        )

        report_period = st.selectbox(
            "📆 Reporting period",
            ["Term 1", "Term 2", "Term 3", "Semester 1", "Semester 2"]
        )

        col_hr, col_sub = st.columns(2)
        with col_hr:
            homeroom_input = st.text_input("🏫 Homeroom (optional)", placeholder="e.g. 7B")
        with col_sub:
            st.caption(f"Subject: **{subject}**")

        topics_covered_input = st.text_area(
            "📘 Topics covered this period",
            placeholder="e.g. Algorithms, Loops, Functions, Intro to Databases",
            help="Used to ground the comment in what was actually taught this term/semester"
        )

        comment_tone = st.selectbox(
            "Tone",
            ["Warm and encouraging", "Formal and professional", "Direct and concise"]
        )
        comment_length = st.selectbox(
            "Length",
            ["Short (1-2 sentences)", "Medium (3-4 sentences)", "Detailed (5+ sentences)"]
        )

        if report_mode == "Single student":
            st.info("✍️ Fill in details for this student's comment")
            student_name = st.text_input("Student name (optional)", placeholder="e.g. Aryan, or leave blank for 'the student'")
            performance_level = st.selectbox(
                "Overall performance level",
                ["Excellent", "Good / Above average", "Satisfactory / Average",
                 "Needs improvement", "Struggling significantly"]
            )
            strengths_input = st.text_area(
                "Strengths (comma separated)",
                placeholder="e.g. problem solving, class participation, neat work"
            )
            improvements_input = st.text_area(
                "Areas to improve (comma separated)",
                placeholder="e.g. time management, showing working, homework completion"
            )
        else:
            st.info("👥 Group your class by performance level. Add student names under each level you need — leave a box empty to skip that level.")
            level_options = ["Excellent", "Good / Above average", "Satisfactory / Average",
                             "Needs improvement", "Struggling significantly"]
            for level in level_options:
                names_text = st.text_area(
                    f"{level} — student names (one per line or comma separated)",
                    placeholder="e.g. Aryan, Maya, Tom",
                    height=70,
                    key=f"batch_{level}"
                )
                if names_text.strip():
                    names = [n.strip() for n in re.split(r"[,\n]", names_text) if n.strip()]
                    if names:
                        batch_groups.append((level, names))

            total_students = sum(len(names) for _, names in batch_groups)
            if total_students:
                st.success(f"📋 {total_students} student(s) across {len(batch_groups)} level(s) ready to generate")

            strengths_input = st.text_area(
                "General strengths to draw from (comma separated)",
                placeholder="e.g. problem solving, class participation, neat work, creativity, teamwork",
                help="The AI will vary which strengths it mentions per student so comments don't feel identical"
            )
            improvements_input = st.text_area(
                "General areas to improve (comma separated)",
                placeholder="e.g. time management, showing working, homework completion, focus"
            )

    # ── Parent Message specific inputs ──
    message_purpose   = ""
    message_details   = ""
    message_platform  = "ClassDojo-style update"

    if task == "💬 Parent Message":
        st.info("💬 What's this message about?")
        message_platform = st.selectbox(
            "Style",
            ["ClassDojo-style update", "Formal email", "Quick reminder/note"]
        )
        message_purpose = st.selectbox(
            "Purpose",
            ["Weekly class update", "Homework reminder", "Behavior/positive note",
             "Upcoming test/event reminder", "General announcement", "Other (describe below)"]
        )
        message_details = st.text_area(
            "Details to include",
            placeholder="e.g. this week we covered fractions, field trip on Friday, please bring permission slip"
        )

    # ── Question/Content customization (all tasks except Lesson Plan) ──
    question_specs = ""
    include_images = False

    if task not in ["📋 Lesson Plan", "🗣️ Report Card Comment", "💬 Parent Message"]:
        with st.expander("🧩 Customize question types & marks", expanded=False):
            include_images = st.checkbox(
                "🖼️ Include image/diagram suggestions",
                help="AI will suggest where diagrams or images should be placed and describe them"
            )

            st.caption("Specify how many of each question type to include (leave 0 to skip):")

            col_a, col_b = st.columns(2)
            with col_a:
                n_mcq        = st.number_input("❓ MCQ", min_value=0, max_value=30, value=0)
                n_fill_blank = st.number_input("✏️ Fill in the blanks", min_value=0, max_value=30, value=0)
                n_match      = st.number_input("🔗 Match the column", min_value=0, max_value=10, value=0)
                n_true_false = st.number_input("☑️ True / False", min_value=0, max_value=30, value=0)
            with col_b:
                n_1marker = st.number_input("1️⃣ 1-mark questions", min_value=0, max_value=30, value=0)
                n_2marker = st.number_input("2️⃣ 2-mark questions", min_value=0, max_value=30, value=0)
                n_4marker = st.number_input("4️⃣ 4-mark questions", min_value=0, max_value=30, value=0)
                n_essay   = st.number_input("📝 Long answer / Essay", min_value=0, max_value=10, value=0)

            spec_parts = []
            if n_mcq:        spec_parts.append(f"{n_mcq} multiple choice questions (MCQ)")
            if n_fill_blank: spec_parts.append(f"{n_fill_blank} fill-in-the-blank questions")
            if n_match:      spec_parts.append(f"{n_match} match-the-column items")
            if n_true_false: spec_parts.append(f"{n_true_false} true/false questions")
            if n_1marker:    spec_parts.append(f"{n_1marker} questions worth 1 mark each")
            if n_2marker:    spec_parts.append(f"{n_2marker} questions worth 2 marks each")
            if n_4marker:    spec_parts.append(f"{n_4marker} questions worth 4 marks each")
            if n_essay:      spec_parts.append(f"{n_essay} long-answer/essay question(s)")

            if spec_parts:
                question_specs = "Include exactly: " + ", ".join(spec_parts) + ". "
            if include_images:
                question_specs += (
                    "Wherever a diagram, illustration, or image would help "
                    "(e.g. for a question or activity), insert a placeholder on its own line "
                    "in EXACTLY this format: [IMAGE: short description of the image to generate] "
                    "— for example [IMAGE: labeled diagram of a plant cell] or "
                    "[IMAGE: a triangle with sides 3, 4 and 5 labeled]. "
                    "Use 1-4 such placeholders depending on relevance. "
                )

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        generate_btn = st.button("⚡ Generate", type="primary", use_container_width=True)
    with col2:
        if st.button("✖ Cancel", use_container_width=True):
            st.session_state.current_output = ""
            st.session_state.current_title  = ""
            st.session_state.edit_mode      = False
            st.session_state.edit_id        = None
            st.rerun()

    if st.button("🆕 New", use_container_width=True):
        st.session_state.current_output = ""
        st.session_state.current_title  = ""
        st.session_state.edit_mode      = False
        st.session_state.edit_id        = None
        st.rerun()

    st.divider()
    st.subheader("📂 My Saved Plans")
    if st.button("🔄 Refresh", use_container_width=True):
        load_plans()

    plans = st.session_state.saved_plans
    if plans:
        for plan in plans:
            col_a, col_b = st.columns([4, 1])
            with col_a:
                label = plan["title"][:25] + "…" if len(plan["title"]) > 25 else plan["title"]
                if st.button(f"📄 {label}", key=f"open_{plan['id']}", use_container_width=True):
                    st.session_state.current_output = plan["content"]
                    st.session_state.current_title  = plan["title"]
                    st.session_state.current_task   = plan.get("task", "")
                    st.session_state.edit_id        = plan["id"]
                    st.session_state.edit_mode      = False
                    st.rerun()
            with col_b:
                if st.button("🗑", key=f"del_{plan['id']}"):
                    delete_plan(plan["id"])
    else:
        st.caption("No saved plans yet. Click 🔄 Refresh after saving, or save your work using 💾 Save to Cloud on the right before leaving.")

# ── Main area ─────────────────────────────
# ── Generate handling ──
if generate_btn:
    is_batch_report = task == "🗣️ Report Card Comment" and report_mode == "Whole class (batch)"
    needs_topics = task not in ["🗣️ Report Card Comment", "💬 Parent Message"]
    missing_subject = not subject
    missing_topics = needs_topics and not topics
    missing_batch = is_batch_report and not batch_groups

    if missing_subject or missing_topics:
        st.warning("Please fill in Subject" + (" and Topics." if needs_topics else "."))
    elif missing_batch:
        st.warning("Please add at least one student name under a performance level.")
    elif is_batch_report:
        total_students = sum(len(names) for _, names in batch_groups)
        title = f"🗣️ Report Card Comments — Whole Class · {subject} · {grade} · {report_period}"
        progress_msg = st.empty()
        comments = []  # list of (name, level, comment_text)
        i = 0
        for level, names in batch_groups:
            for name in names:
                i += 1
                progress_msg.info(f"✍️ Writing comment {i}/{total_students}: {name}...")
                try:
                    prompt = build_prompt(
                        task, subject, grade, topics,
                        lessons_per_week, duration, curriculum,
                        activity_duration, test_duration,
                        question_specs, name, level,
                        strengths_input, improvements_input, comment_tone,
                        comment_length, message_platform, message_purpose,
                        message_details, report_period, topics_covered_input,
                        homeroom_input
                    )
                    comment_text = generate(prompt)
                    comments.append((name, level, comment_text))
                except Exception as e:
                    comments.append((name, level, f"[Error generating comment: {e}]"))
        progress_msg.empty()

        # Combine into one readable document
        homeroom_header = f" · Homeroom {homeroom_input.strip()}" if homeroom_input.strip() else ""
        combined_parts = [f"# Report Card Comments — {subject}, {grade}{homeroom_header}\n{report_period}\n"]
        for level, names in batch_groups:
            level_comments = [(n, c) for n, l, c in comments if l == level]
            if level_comments:
                combined_parts.append(f"\n## {level}\n")
                for name, comment_text in level_comments:
                    combined_parts.append(f"**{name}:**\n{comment_text}\n")
        combined_output = "\n".join(combined_parts)

        st.session_state.current_output = combined_output
        st.session_state.current_title  = title
        st.session_state.current_task   = task
        st.session_state.edit_mode      = False
        st.session_state.edit_id        = None
        st.session_state.generated_images = {}
        st.success(f"✅ Generated {len(comments)} comments for the whole class!")
    else:
        if task == "🗣️ Report Card Comment":
            title = f"{task} — {student_name.strip() or 'Student'} · {subject} · {grade} · {report_period}"
        elif task == "💬 Parent Message":
            title = f"{task} — {message_purpose} · {subject} · {grade}"
        else:
            title = f"{task} — {curriculum} · {subject} · {grade}"
        with st.spinner("Generating... ⏳"):
            try:
                result = generate(build_prompt(
                    task, subject, grade, topics,
                    lessons_per_week, duration, curriculum,
                    activity_duration, test_duration,
                    question_specs, student_name, performance_level,
                    strengths_input, improvements_input, comment_tone,
                    comment_length, message_platform, message_purpose,
                    message_details, report_period, topics_covered_input,
                    homeroom_input
                ))
                st.session_state.current_output = result
                st.session_state.current_title  = title
                st.session_state.current_task   = task
                st.session_state.edit_mode      = False
                st.session_state.edit_id        = None

                # Generate actual images for any [IMAGE: ...] placeholders
                img_prompts = extract_image_prompts(result)
                generated_images = {}
                debug_msgs = []
                if img_prompts:
                    progress_msg = st.empty()
                    for i, desc in enumerate(img_prompts, 1):
                        progress_msg.info(f"🎨 Generating image {i}/{len(img_prompts)}: {desc[:50]}...")
                        img_bytes, debug_info = generate_image(desc)
                        if img_bytes:
                            generated_images[desc] = img_bytes
                        else:
                            debug_msgs.append(f"'{desc[:40]}...' → {debug_info}")
                    progress_msg.empty()
                    n_ok = len(generated_images)
                    if n_ok < len(img_prompts):
                        st.warning(f"⚠️ Generated {n_ok}/{len(img_prompts)} images. "
                                   "Remaining show as suggestions with search links.")
                        with st.expander("🔧 Debug info (why images failed)"):
                            for msg in debug_msgs:
                                st.code(msg)
                st.session_state.generated_images = generated_images
            except Exception as e:
                st.error(f"Error: {e}")

# ── Main content display ──
if st.session_state.current_output:
    st.subheader(st.session_state.current_title)

    images = st.session_state.get("generated_images", {})

    # Action toolbar
    action_cols = st.columns(6)
    with action_cols[0]:
        if st.button("💾 Save", use_container_width=True):
            save_plan(
                st.session_state.current_title,
                st.session_state.current_output,
                st.session_state.current_task,
                curriculum, subject, grade
            )
    with action_cols[1]:
        if st.button("✏️ Edit", use_container_width=True):
            st.session_state.edit_mode = True
            st.rerun()

    word_buf = export_word(st.session_state.current_title, st.session_state.current_output, images)
    with action_cols[2]:
        st.download_button(
            "⬇️ Word", data=word_buf,
            file_name=f"{st.session_state.current_title[:40]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    if "Lesson Plan" in st.session_state.current_task:
        xl_buf = export_excel(st.session_state.current_title, st.session_state.current_output)
        with action_cols[3]:
            st.download_button(
                "📊 Excel", data=xl_buf,
                file_name=f"{st.session_state.current_title[:40]}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

    with action_cols[4]:
        st.download_button(
            "📄 Text", data=st.session_state.current_output,
            file_name=f"{st.session_state.current_title[:40]}.txt",
            mime="text/plain", use_container_width=True
        )

    with action_cols[5]:
        if st.button("🗑 Clear", use_container_width=True):
            st.session_state.current_output = ""
            st.session_state.current_title  = ""
            st.session_state.edit_mode      = False
            st.session_state.edit_id        = None
            st.session_state.generated_images = {}
            st.rerun()

    st.divider()

    if st.session_state.edit_mode:
        edited = st.text_area("✏️ Edit:", value=st.session_state.current_output, height=500)
        c1, c2 = st.columns(2)
        with c1:
            if st.button("✅ Save edits", type="primary"):
                st.session_state.current_output = edited
                if st.session_state.edit_id:
                    save_plan(
                        st.session_state.current_title, edited,
                        st.session_state.current_task, "", "", ""
                    )
                st.session_state.edit_mode = False
                st.rerun()
        with c2:
            if st.button("❌ Cancel edit"):
                st.session_state.edit_mode = False
                st.rerun()
    else:
        content = st.session_state.current_output
        images  = st.session_state.get("generated_images", {})

        # Split content on [IMAGE: ...] placeholders and render inline
        parts = re.split(r"(\[IMAGE:\s*.*?\])", content, flags=re.IGNORECASE)
        for part in parts:
            m = re.match(r"\[IMAGE:\s*(.*?)\]", part, re.IGNORECASE)
            if m:
                desc = m.group(1).strip()
                img_bytes = images.get(desc)
                if img_bytes:
                    st.image(img_bytes, caption=desc, width=400)
                else:
                    st.info(f"🖼️ Image suggestion: *{desc}*")
                    st.markdown(f"[🔎 Search reference images]({reference_image_search_url(desc)})")
            elif part.strip():
                st.markdown(part)
else:
    st.info("👈 Fill in your details on the left and click ⚡ Generate")